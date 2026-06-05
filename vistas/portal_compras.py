"""
portal_compras.py — Vistas del módulo de compras de entradas.

Implementa el flujo completo de compra de tickets con integración a Mercado Pago:

FLUJO:
1. seleccion_asientos_view: Muestra la sala con los asientos disponibles/ocupados.
2. procesar_pago_view: Verifica asientos, guarda en sesión, crea preferencia MP,
   retorna la URL de pago para que el frontend abra en nueva pestaña.
3. verificar_pago_view: El usuario vuelve y hace clic en "Ya pagué".
   Busca el pago en la API de MP usando external_reference.
   Si aprobado → crea Venta + Tickets + QR.
4. retorno_mercadopago_view: Procesa el retorno si MP logra redirigir (backup).
5. ver_tickets_view: Muestra los tickets generados con sus QR.
"""

import json
import uuid
import qrcode
import zipfile
from io import BytesIO
from decimal import Decimal

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.db import transaction, IntegrityError
from django.contrib import messages
from django.http import JsonResponse, HttpResponse, Http404
from django.views.decorators.http import require_POST
from django.core.files.base import ContentFile

from sistema_cine.models import Funcion, Asiento, Venta, Ticket


# ─── VISTA 1: SELECCIÓN DE ASIENTOS ──────────────────────────────────────────

@login_required
def seleccion_asientos_view(request, funcion_id):
    """
    Renderiza la sala y sus asientos para una función específica.
    """
    funcion = get_object_or_404(Funcion, id=funcion_id)
    sala = funcion.sala
    asientos = sala.asientos.filter(estado_asiento='activo')

    asientos_ocupados_ids = Ticket.objects.filter(
        funcion=funcion
    ).values_list('asiento_id', flat=True)

    layout = sala.layout_config if sala.layout_config else {}

    return render(request, 'portal_cliente/seleccion_asientos.html', {
        'funcion': funcion,
        'sala': sala,
        'layout': json.dumps(layout),
        'asientos_ocupados': list(asientos_ocupados_ids),
        'asientos_totales': asientos
    })


# ─── VISTA 2: PROCESAR PAGO (CREAR PREFERENCIA EN MP) ────────────────────────

@login_required
@require_POST
def procesar_pago_view(request, funcion_id):
    """
    Crea una preferencia de pago en Mercado Pago.
    Retorna la URL de pago para que el frontend abra en nueva pestaña.
    """
    funcion = get_object_or_404(Funcion, id=funcion_id)

    try:
        data = json.loads(request.body)
        asientos_ids = data.get('asientos_ids', [])

        if not asientos_ids:
            return JsonResponse({'error': 'No seleccionaste ningún asiento.'}, status=400)

        # Verificar que los asientos existen
        asientos = Asiento.objects.filter(id__in=asientos_ids)
        if asientos.count() != len(asientos_ids):
            return JsonResponse({'error': 'Algunos asientos no son válidos.'}, status=400)

        # ─── BARRERA 1: Verificar disponibilidad ──────────────────────────────
        asientos_ocupados = Ticket.objects.filter(
            funcion=funcion,
            asiento_id__in=asientos_ids
        ).values_list('asiento_id', flat=True)

        if asientos_ocupados:
            return JsonResponse({
                'error': 'Algunos asientos ya fueron comprados. Recargá la página.'
            }, status=409)

        monto_total = len(asientos_ids) * funcion.precio_entrada

        # Generar referencia única para vincular este pago con nuestra sesión
        referencia = f"cinefsa-{uuid.uuid4().hex[:12]}"

        # Guardar en sesión
        request.session['compra_pendiente'] = {
            'funcion_id': funcion_id,
            'asientos_ids': asientos_ids,
            'monto_total': str(monto_total),
            'referencia': referencia,
        }

        # Crear preferencia en Mercado Pago
        from servicios.mercadopago_service import ServicioMercadoPago
        mp = ServicioMercadoPago()
        resultado = mp.crear_preferencia(
            funcion=funcion,
            cantidad_asientos=len(asientos_ids),
            monto_total=monto_total,
            external_reference=referencia,
            request=request,
        )

        return JsonResponse({
            'init_point': resultado['init_point'],
            'referencia': referencia,
        })

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


# ─── VISTA 3: VERIFICAR PAGO (el usuario hace clic en "Ya pagué") ────────────

@login_required
def verificar_pago_view(request):
    """
    Busca en la API de Mercado Pago si existe un pago aprobado para
    la referencia guardada en la sesión.
    
    Se llama por AJAX cuando el usuario hace clic en "Ya pagué".
    """
    compra = request.session.get('compra_pendiente')
    if not compra:
        return JsonResponse({'status': 'error', 'mensaje': 'No hay compra en proceso.'}, status=400)

    referencia = compra.get('referencia')
    if not referencia:
        return JsonResponse({'status': 'error', 'mensaje': 'Referencia no encontrada.'}, status=400)

    # Buscar el pago en la API de MP usando la referencia
    from servicios.mercadopago_service import ServicioMercadoPago
    mp = ServicioMercadoPago()
    pago = mp.buscar_pago_por_referencia(referencia)

    if not pago:
        return JsonResponse({
            'status': 'not_found',
            'mensaje': 'Todavía no detectamos tu pago. Si ya pagaste, esperá unos segundos e intentá de nuevo.',
        })

    if pago['status'] == 'approved':
        # ─── PAGO APROBADO: Crear Venta + Tickets + QR ───────────────────
        resultado = _confirmar_compra(request, compra)
        return JsonResponse(resultado)

    elif pago['status'] in ('pending', 'in_process'):
        return JsonResponse({
            'status': 'pending',
            'mensaje': 'Tu pago está siendo procesado. Esperá unos momentos.',
        })

    else:
        # Rechazado u otro error
        if 'compra_pendiente' in request.session:
            del request.session['compra_pendiente']
        return JsonResponse({
            'status': 'rejected',
            'mensaje': f'Tu pago fue rechazado. Detalle: {pago.get("status_detail", "desconocido")}',
        })


# ─── VISTA 4: RETORNO DESDE MERCADO PAGO (backup si MP logra redirigir) ──────

@login_required
def retorno_mercadopago_view(request):
    """
    Backup: Si Mercado Pago logra redirigir al usuario (en producción con dominio real),
    esta vista procesa el retorno.
    """
    payment_id = request.GET.get('payment_id')
    compra = request.session.get('compra_pendiente')

    if not compra:
        messages.warning(request, 'No hay una compra en proceso.')
        return redirect('inicio')

    if not payment_id:
        if 'compra_pendiente' in request.session:
            del request.session['compra_pendiente']
        return render(request, 'portal_cliente/resultado_pago.html', {
            'resultado': 'cancelled',
        })

    # Verificar pago con la API
    from servicios.mercadopago_service import ServicioMercadoPago
    mp = ServicioMercadoPago()
    pago = mp.verificar_pago(payment_id)

    if pago['status'] == 'approved':
        resultado = _confirmar_compra(request, compra)
        if resultado['status'] == 'approved':
            return redirect('compras:ver_tickets', venta_id=resultado['venta_id'])
        else:
            return render(request, 'portal_cliente/resultado_pago.html', {
                'resultado': 'seats_taken',
            })

    elif pago['status'] in ('pending', 'in_process'):
        return render(request, 'portal_cliente/resultado_pago.html', {
            'resultado': 'pending',
        })

    else:
        if 'compra_pendiente' in request.session:
            del request.session['compra_pendiente']
        return render(request, 'portal_cliente/resultado_pago.html', {
            'resultado': 'rejected',
            'detalle': pago.get('status_detail', ''),
        })


# ─── FUNCIÓN INTERNA: CONFIRMAR COMPRA ───────────────────────────────────────

def _confirmar_compra(request, compra):
    """
    Lógica compartida para confirmar una compra una vez que el pago fue aprobado.
    Verifica asientos libres, crea Venta + Tickets + QR.
    
    Returns:
        dict con 'status' y datos adicionales.
    """
    funcion = get_object_or_404(Funcion, id=compra['funcion_id'])
    asientos_ids = compra['asientos_ids']
    monto_total = Decimal(compra['monto_total'])

    try:
        with transaction.atomic():
            # ─── BARRERA 2: Verificar asientos libres dentro de transacción ───
            asientos_ocupados = list(
                Ticket.objects.filter(
                    funcion=funcion,
                    asiento_id__in=asientos_ids
                ).select_for_update().values_list('asiento_id', flat=True)
            )

            if asientos_ocupados:
                raise IntegrityError("Asientos ocupados")

            # Crear Venta
            venta = Venta.objects.create(
                usuario=request.user,
                monto_total=monto_total,
                metodo_pago='mercado_pago',
                estado_pago='aprobado',
            )

            # Crear Tickets
            asientos = Asiento.objects.filter(id__in=asientos_ids)
            tickets_creados = []
            for asiento in asientos:
                ticket = Ticket.objects.create(
                    venta=venta,
                    funcion=funcion,
                    asiento=asiento,
                    estado_uso='pendiente',
                )
                tickets_creados.append(ticket)

        # Generar QR fuera del atomic
        for ticket in tickets_creados:
            _generar_qr_ticket(ticket)

        # Limpiar sesión
        if 'compra_pendiente' in request.session:
            del request.session['compra_pendiente']

        return {
            'status': 'approved',
            'mensaje': '¡Pago exitoso! Tus tickets están listos.',
            'venta_id': venta.id_venta,
            'redirect_url': f'/compras/venta/{venta.id_venta}/tickets/',
        }

    except IntegrityError:
        if 'compra_pendiente' in request.session:
            del request.session['compra_pendiente']
        return {
            'status': 'seats_taken',
            'mensaje': 'Los asientos fueron comprados por otro usuario mientras pagabas.',
        }


# ─── UTILIDAD: GENERAR QR ─────────────────────────────────────────────────────

def _generar_qr_ticket(ticket):
    """Genera una imagen QR para un ticket."""
    contenido_qr = f"CINEFSA-TICKET-{ticket.codigo_qr}"

    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_H,
        box_size=10,
        border=4,
    )
    qr.add_data(contenido_qr)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")

    buffer = BytesIO()
    img.save(buffer, format='PNG')
    buffer.seek(0)

    nombre_archivo = f"ticket_{ticket.id_ticket}_{ticket.codigo_qr}.png"
    ticket.imagen_qr.save(nombre_archivo, ContentFile(buffer.read()), save=True)


# ─── VISTA 5: VER TICKETS ────────────────────────────────────────────────────

@login_required
def ver_tickets_view(request, venta_id):
    """Muestra los tickets generados de una venta con sus códigos QR."""
    venta = get_object_or_404(Venta, id_venta=venta_id, usuario=request.user)
    tickets = venta.tickets.select_related('funcion__pelicula', 'asiento', 'funcion__sala')

    return render(request, 'portal_cliente/ver_tickets.html', {
        'venta': venta,
        'tickets': tickets,
        'titulo_pagina': f'Tickets - Venta #{venta.id_venta}',
    })


# ─── VISTA 6: DESCARGAR QR INDIVIDUAL ────────────────────────────────────────

@login_required
def descargar_qr_ticket(request, ticket_id):
    """Descarga la imagen QR de un ticket individual como PNG."""
    ticket = get_object_or_404(Ticket, id_ticket=ticket_id, venta__usuario=request.user)

    if not ticket.imagen_qr:
        raise Http404("Este ticket no tiene QR generado.")

    response = HttpResponse(ticket.imagen_qr.read(), content_type='image/png')
    nombre = f"CineFSA_Ticket_{ticket.id_ticket}_{ticket.asiento.fila}{ticket.asiento.numero}.png"
    response['Content-Disposition'] = f'attachment; filename="{nombre}"'
    return response


# ─── VISTA 7: DESCARGAR ZIP DE QR ────────────────────────────────────────────

@login_required
def descargar_qr_venta(request, venta_id):
    """Descarga todos los QR de una venta en un archivo ZIP."""
    venta = get_object_or_404(Venta, id_venta=venta_id, usuario=request.user)
    tickets = venta.tickets.select_related('funcion__pelicula', 'asiento')

    buffer = BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for ticket in tickets:
            if ticket.imagen_qr:
                nombre = f"Ticket_{ticket.id_ticket}_{ticket.asiento.fila}{ticket.asiento.numero}.png"
                zf.writestr(nombre, ticket.imagen_qr.read())

    buffer.seek(0)

    response = HttpResponse(buffer.read(), content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="CineFSA_Venta_{venta.id_venta}_QRs.zip"'
    return response


# ─── VISTA 8: DESCARGAR TICKET COMPLETO COMO WORD (.docx) ────────────────────

@login_required
def descargar_ticket_word(request, ticket_id):
    """
    Genera y descarga un documento Word (.docx) con el ticket completo.

    Usa una consulta SQL directa (raw) para obtener todos los datos del ticket
    desde la base de datos, y luego construye el documento con python-docx.

    Datos incluidos en el ticket:
    - Película (título)
    - Sala (nombre)
    - Asiento (fila + número + tipo)
    - Fecha y hora de la función
    - Precio de la entrada
    - Código QR (imagen embebida en el Word)
    - Estado del ticket (pendiente / validado)
    """
    from django.db import connection
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import os

    # ── Consulta SQL directa para obtener TODOS los datos del ticket ─────
    sql = """
        SELECT
            t.id_ticket,
            t.codigo_qr,
            t.estado_uso,
            t.imagen_qr,
            p.titulo        AS pelicula,
            s.nombre_sala   AS sala,
            a.fila           AS asiento_fila,
            a.numero         AS asiento_numero,
            a.tipo_asiento   AS tipo_asiento,
            f.fecha          AS fecha_funcion,
            f.hora_inicio    AS hora_funcion,
            f.precio_entrada AS precio,
            v.id_venta,
            v.fecha_hora_transaccion,
            v.monto_total,
            v.usuario_id
        FROM ticket t
        INNER JOIN funcion f  ON t.funcion_id = f.id
        INNER JOIN pelicula p ON f.pelicula_id = p.id
        INNER JOIN sala s     ON f.sala_id = s.id
        INNER JOIN asiento a  ON t.asiento_id = a.id
        INNER JOIN venta v    ON t.venta_id = v.id_venta
        WHERE t.id_ticket = %s
          AND v.usuario_id = %s
    """

    with connection.cursor() as cursor:
        cursor.execute(sql, [ticket_id, request.user.id])
        columnas = [col[0] for col in cursor.description]
        row = cursor.fetchone()

    if not row:
        raise Http404("Ticket no encontrado.")

    # Mapear columnas a un diccionario
    datos = dict(zip(columnas, row))

    # Mapear tipo de asiento
    tipos_asiento = {'general': 'General', 'vip': 'VIP', 'discapacitado': 'Discapacitado'}
    tipo_display = tipos_asiento.get(datos['tipo_asiento'], datos['tipo_asiento'])

    # Mapear estado
    estados = {'pendiente': 'Pendiente de Validación', 'validado': 'Ingreso Validado'}
    estado_display = estados.get(datos['estado_uso'], datos['estado_uso'])

    # ── Construir el documento Word ──────────────────────────────────────
    doc = Document()

    # Configurar márgenes más estrechos para que parezca un ticket
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ── Encabezado: CINEFSA ──────────────────────────────────────────────
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run('CINEFSA')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run('ENTRADA DE CINE')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.bold = True

    # Línea separadora
    doc.add_paragraph('━' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Título de la película ────────────────────────────────────────────
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_titulo.add_run(datos['pelicula'])
    run.font.size = Pt(20)
    run.font.bold = True

    doc.add_paragraph()  # Espacio

    # ── Tabla con datos del ticket (sin bordes, limpia) ───────────────────
    tabla = doc.add_table(rows=4, cols=2)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    campos = [
        ('SALA',    datos['sala']),
        ('FECHA',   str(datos['fecha_funcion'].strftime('%d/%m/%Y') if hasattr(datos['fecha_funcion'], 'strftime') else datos['fecha_funcion'])),
        ('HORARIO', str(datos['hora_funcion'].strftime('%H:%M') if hasattr(datos['hora_funcion'], 'strftime') else datos['hora_funcion']) + ' hs'),
        ('PRECIO',  f"${datos['precio']}"),
    ]

    for i, (label, valor) in enumerate(campos):
        # Celda label
        cell_label = tabla.cell(i, 0)
        p = cell_label.paragraphs[0]
        run = p.add_run(label)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.font.bold = True

        # Celda valor
        cell_valor = tabla.cell(i, 1)
        p = cell_valor.paragraphs[0]
        run = p.add_run(valor)
        run.font.size = Pt(13)
        run.font.bold = True

    doc.add_paragraph()  # Espacio

    # ── Asiento destacado ────────────────────────────────────────────────
    p_asiento = doc.add_paragraph()
    p_asiento.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_asiento.add_run(f'BUTACA {datos["asiento_fila"]}{datos["asiento_numero"]}')
    run.font.size = Pt(18)
    run.font.bold = True

    p_tipo = doc.add_paragraph()
    p_tipo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tipo.add_run(tipo_display)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Línea separadora
    doc.add_paragraph('━' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Código QR (imagen embebida) ──────────────────────────────────────
    if datos['imagen_qr']:
        from django.conf import settings
        qr_path = os.path.join(settings.MEDIA_ROOT, datos['imagen_qr'])
        if os.path.exists(qr_path):
            p_qr = doc.add_paragraph()
            p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_qr.add_run()
            run.add_picture(qr_path, width=Inches(2))

    # Código UUID
    p_codigo = doc.add_paragraph()
    p_codigo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_codigo.add_run(str(datos['codigo_qr']))
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(150, 150, 150)

    # Estado
    p_estado = doc.add_paragraph()
    p_estado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_estado.add_run(estado_display.upper())
    run.font.size = Pt(9)
    run.font.bold = True
    if datos['estado_uso'] == 'validado':
        run.font.color.rgb = RGBColor(6, 95, 70)
    else:
        run.font.color.rgb = RGBColor(146, 64, 14)

    # ── Pie: info de la venta ────────────────────────────────────────────
    doc.add_paragraph('━' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_pie.add_run(f'Venta #{datos["id_venta"]}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(130, 130, 130)

    # ── Guardar y devolver como respuesta HTTP ───────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    pelicula_safe = datos['pelicula'].replace(' ', '_')[:30]
    asiento_str = f"{datos['asiento_fila']}{datos['asiento_numero']}"
    nombre = f"CineFSA_Ticket_{pelicula_safe}_{asiento_str}.docx"

    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{nombre}"'
    return response
